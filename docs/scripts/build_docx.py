#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_docx.py — Sinh báo cáo đồ án (.docx) từ tài liệu Markdown nguồn.

Đọc  docs/BaoCao_QuanLyLuong.md  và "extract" ra một file Word chuẩn đồ án:
trang bìa, mục lục tự động (TOC), tiêu đề nhiều cấp, đoạn văn căn đều, bảng kẻ
khung, danh sách, trích dẫn và hình ảnh (kèm chú thích).

Cách dùng:
    python docs/scripts/build_docx.py
    python docs/scripts/build_docx.py <file.md> <file_ra.docx>

Yêu cầu: pip install python-docx
"""
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

# ---------------------------------------------------------------------------
#  Thông tin trang bìa — chỉnh sửa cho phù hợp đồ án của bạn.
# ---------------------------------------------------------------------------
COVER = {
    "truong": "TRƯỜNG ĐẠI HỌC ………………",
    "khoa": "KHOA CÔNG NGHỆ THÔNG TIN",
    "loai": "BÁO CÁO ĐỒ ÁN",
    "tieu_de": "PHÂN HỆ QUẢN LÝ LƯƠNG & CHẤM CÔNG\nHỆ THỐNG POS CHUỖI CỬA HÀNG TIỆN LỢI MINIMART",
    "sinh_vien": "Hồ Huy Hoàng",
    "gvhd": "………………………………",
    "nam": "2026",
}

BASE_FONT = "Times New Roman"
BASE_SIZE = 13
HEADING_COLOR = RGBColor(0x0B, 0x5A, 0x3C)   # xanh lá đậm — đồng bộ nhận diện MiniMart

# ---------------------------------------------------------------------------
#  Tiện ích định dạng
# ---------------------------------------------------------------------------
def set_base_styles(doc):
    """Đặt font nền Times New Roman 13, giãn dòng 1.5; restyle các cấp tiêu đề."""
    normal = doc.styles["Normal"]
    normal.font.name = BASE_FONT
    normal.font.size = Pt(BASE_SIZE)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)

    sizes = {"Heading 1": 16, "Heading 2": 14, "Heading 3": 13, "Heading 4": 13}
    for name, size in sizes.items():
        st = doc.styles[name]
        st.font.name = BASE_FONT
        st.element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = HEADING_COLOR
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(4)


def add_runs(paragraph, text):
    """Thêm text vào paragraph, hiểu **đậm**, *nghiêng*, `mã`."""
    token = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
    for part in token.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(BASE_SIZE - 1)
        else:
            paragraph.add_run(part)


def add_toc(doc):
    """Chèn trường MỤC LỤC tự động (Word cập nhật khi mở: Update Field)."""
    p = doc.add_paragraph()
    run = p.add_run()
    for kind, txt in (("begin", None), ("instr", 'TOC \\o "1-3" \\h \\z \\u'), ("sep", None)):
        if kind == "instr":
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = txt
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), "begin" if kind == "begin" else "separate")
        run._r.append(el)
    hint = p.add_run("Mục lục sẽ hiển thị khi mở file Word và cập nhật trường (chuột phải → Update Field).")
    hint.italic = True
    hint.font.size = Pt(BASE_SIZE - 2)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    p.add_run()._r.append(end)


def build_cover(doc):
    """Trang bìa căn giữa."""
    def center(text, size, bold=False, space_before=0, space_after=10, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        for i, line in enumerate(text.split("\n")):
            if i:
                p.add_run().add_break()
            r = p.add_run(line)
            r.bold = bold
            r.font.size = Pt(size)
            r.font.name = BASE_FONT
            if color:
                r.font.color.rgb = color
        return p

    center(COVER["truong"], 14, bold=True, space_before=18)
    center(COVER["khoa"], 13, bold=True, space_after=60)
    center(COVER["loai"], 18, bold=True, color=HEADING_COLOR, space_after=24)
    center(COVER["tieu_de"], 20, bold=True, color=HEADING_COLOR, space_after=70)
    center(f"Sinh viên thực hiện: {COVER['sinh_vien']}", 13, space_after=6)
    center(f"Giảng viên hướng dẫn: {COVER['gvhd']}", 13, space_after=60)
    center(COVER["nam"], 13, bold=True)
    doc.add_page_break()
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("MỤC LỤC")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = HEADING_COLOR
    add_toc(doc)
    doc.add_page_break()


def add_table(doc, rows):
    """rows: list các list ô (đã tách theo '|'); dòng đầu là tiêu đề."""
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, cells in enumerate(rows):
        row = table.add_row().cells
        for ci in range(cols):
            text = cells[ci] if ci < len(cells) else ""
            para = row[ci].paragraphs[0]
            add_runs(para, text)
            for run in para.runs:
                run.font.size = Pt(BASE_SIZE - 1)
                if ri == 0:
                    run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_image(doc, alt, path, base_dir):
    full = os.path.normpath(os.path.join(base_dir, path))
    if not os.path.exists(full):
        doc.add_paragraph(f"[Thiếu hình: {path}]").italic = True
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(full, width=Inches(6.0))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run("Hình: " + alt)
    r.italic = True
    r.font.size = Pt(BASE_SIZE - 2)


# ---------------------------------------------------------------------------
#  Bộ phân tích Markdown (tập con dùng trong tài liệu)
# ---------------------------------------------------------------------------
def render(doc, md_text, base_dir):
    lines = md_text.splitlines()
    i, n = 0, len(lines)
    img_re = re.compile(r"^!\[(.*?)\]\((.*?)\)\s*$")
    while i < n:
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Tiêu đề
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2).strip(), level=level)
            i += 1
            continue

        # Hình ảnh
        m = img_re.match(stripped)
        if m:
            add_image(doc, m.group(1), m.group(2), base_dir)
            i += 1
            continue

        # Bảng (khối các dòng chứa '|')
        if stripped.startswith("|") and "|" in stripped[1:]:
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            rows = []
            for b in block:
                if re.match(r"^\|[\s:\-\|]+\|?$", b):   # bỏ dòng phân cách ---
                    continue
                cells = [c.strip() for c in b.strip().strip("|").split("|")]
                rows.append(cells)
            if rows:
                add_table(doc, rows)
            continue

        # Trích dẫn
        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            add_runs(p, stripped.lstrip("> ").strip())
            i += 1
            continue

        # Danh sách có thứ tự
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        # Danh sách gạch đầu dòng
        if stripped.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, stripped[2:])
            i += 1
            continue

        # Đường kẻ ngang → bỏ qua
        if stripped == "---":
            i += 1
            continue

        # Đoạn văn thường (căn đều)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(p, stripped)
        i += 1


def main():
    here = os.path.dirname(os.path.abspath(__file__))
    docs_dir = os.path.abspath(os.path.join(here, ".."))
    md_path = sys.argv[1] if len(sys.argv) > 1 else os.path.join(docs_dir, "BaoCao_QuanLyLuong.md")
    out_path = sys.argv[2] if len(sys.argv) > 2 else os.path.join(docs_dir, "BaoCao_QuanLyLuong.docx")

    with open(md_path, encoding="utf-8") as f:
        md_text = f.read()

    doc = Document()
    set_base_styles(doc)
    build_cover(doc)
    render(doc, md_text, base_dir=os.path.dirname(md_path))
    doc.save(out_path)
    print(f"Đã tạo: {out_path}")


if __name__ == "__main__":
    main()
